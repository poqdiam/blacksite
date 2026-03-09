#!/usr/bin/env python3
"""
BLACKSITE SOP Document Generator
Produces one DOCX per workflow + an index DOCX from screenshots in workflows/screenshots/.
Output: workflows/*.docx
"""
from __future__ import annotations
import os
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Paths ───────────────────────────────────────────────────────────────────

SHOT_DIR  = Path(__file__).parent.parent / "workflows" / "screenshots"
OUT_DIR   = Path(__file__).parent.parent / "workflows"
TODAY     = date.today().strftime("%Y-%m-%d")

OUT_DIR.mkdir(parents=True, exist_ok=True)

# ─── Brand constants ──────────────────────────────────────────────────────────

BRAND_BLUE  = RGBColor(0x1a, 0x1a, 0x2e)   # dark navy
BRAND_CYAN  = RGBColor(0x00, 0xd9, 0xf5)   # BLACKSITE cyan accent
BRAND_WHITE = RGBColor(0xff, 0xff, 0xff)
BRAND_GRAY  = RGBColor(0x44, 0x47, 0x5a)
TEXT_DARK   = RGBColor(0x1e, 0x1e, 0x2e)

SYSTEM_NAME = "BLACKSITE GRC Platform"
SYSTEM_ID   = "THKR-BSV-001"
DOC_VERSION = "1.0"

# ─── Workflow definitions ─────────────────────────────────────────────────────
# Each workflow: slug, title, role(s), description, steps list
# Each step: (number, heading, body_text, screenshot_label_or_None)

WORKFLOWS = [

  # ── 01 SESSION AUTH ────────────────────────────────────────────────────────
  {
    "slug":  "session-auth",
    "title": "Session Authentication and Access",
    "roles": ["All roles"],
    "desc":  (
      "BLACKSITE uses a reverse proxy (Authelia) to authenticate users before any "
      "page is served. The Remote-User header carries the authenticated username. "
      "This SOP covers how sessions are established, the idle timeout policy, "
      "and how to verify system health."
    ),
    "prereqs": [
      "Your username must exist in the BLACKSITE user database.",
      "You must be authenticated through the Authelia SSO portal at your organization's URL.",
      "Your browser must accept first-party cookies from the BLACKSITE domain.",
    ],
    "steps": [
      (1,  "Access the BLACKSITE URL",
           "Navigate to your organization's BLACKSITE URL. Authelia will intercept the request "
           "and verify your session. If you are not logged in, you will be redirected to the "
           "Authelia login page. Complete SSO authentication there first.",
           "no-header-401"),
      (2,  "Confirm dashboard loads",
           "After successful authentication, Authelia injects the Remote-User header and forwards "
           "the request. BLACKSITE reads this header and identifies your account. You will be "
           "redirected to your role-appropriate dashboard automatically.",
           "dashboard-redirect"),
      (3,  "Verify system health (optional)",
           "Navigate to /health to confirm the application and database are operational. "
           "This endpoint returns a JSON payload with status 'ok' and the current timestamp. "
           "Use this endpoint to confirm connectivity if the dashboard does not load.",
           "health-endpoint"),
      (4,  "Session idle timeout behavior",
           "BLACKSITE enforces a 15-minute idle session timeout. If you remain inactive, "
           "a toast notification will appear warning you. Clicking anywhere on the page "
           "resets the timer. The heartbeat API is called automatically every 5 minutes "
           "while the page is open.",
           "session-timeout-toast"),
      (5,  "Sign out",
           "Navigate to /logout or click your username in the top-right navigation and select "
           "Sign Out. BLACKSITE clears its session state. You will be redirected to the "
           "Authelia logout page. Note: Authelia may keep your SSO session active.",
           "logout-redirect"),
      (6,  "Check API version",
           "Navigate to /api/version to confirm the running build stamp, version string, "
           "and deployment timestamp. Use this when troubleshooting version-specific issues "
           "or verifying a new deployment went live.",
           "api-version"),
    ],
    "failures": [
      ("401 Unauthorized page", "Your Authelia session expired. Return to the Authelia portal and re-authenticate."),
      ("Dashboard blank / spinner stuck", "The database may be unavailable. Check /health. Contact the system administrator."),
      ("Session logs out immediately", "Your account may be frozen. Contact an admin to verify account status at /admin/users."),
    ],
  },

  # ── 02 DASHBOARD ACCESS ────────────────────────────────────────────────────
  {
    "slug":  "dashboard-access",
    "title": "Role Dashboards and Navigation",
    "roles": ["All roles"],
    "desc":  (
      "BLACKSITE serves a role-appropriate dashboard to each user upon login. "
      "Platform roles (admin, principal, executive, manager, analyst) and system roles "
      "(ISSO, ISSM, SCA, AO, CISO, etc.) each receive a tailored landing page with "
      "relevant metrics, pending actions, and quick links. This SOP covers the main "
      "dashboard variants and the global navigation sidebar."
    ),
    "prereqs": [
      "Authenticated session established (see SOP: Session Authentication).",
      "Account exists in BLACKSITE with a role assigned.",
    ],
    "steps": [
      (1,  "Admin dashboard",
           "Admin users see a global summary: all systems, active users, recent audit events, "
           "and pending submissions. Key metric cards appear across the top. The left sidebar "
           "shows admin-only links (Users, Audit, Ingest, System Settings).",
           "admin-dashboard"),
      (2,  "ISSO dashboard",
           "ISSO users see their assigned systems with control coverage bars, open POA&M counts, "
           "overdue items, and a daily tasks widget. The daily quiz prompt appears if today's quiz "
           "is not yet completed.",
           "isso-dashboard"),
      (3,  "ISSM dashboard",
           "ISSMs see a portfolio view of all systems under their oversight. Each system row "
           "shows its ISSO, open POA&M count, overdue count, ATO expiry, and rotation day. "
           "Click a system row to open its detail.",
           "issm-dashboard"),
      (4,  "SCA / CISO dashboard",
           "SCA users see their workspace queue — assessments pending review and completed "
           "assessment reports. CISO users see an executive posture summary with aggregate "
           "risk metrics across the portfolio.",
           "ciso-dashboard"),
      (5,  "AO decisions queue",
           "Authorizing Officials see their pending authorization decisions, risk acceptances "
           "awaiting signature, and recent ATO decisions. Each row links to the system's "
           "ATO document matrix.",
           "ao-decisions"),
      (6,  "SCA workspace",
           "Security Control Assessors access their personal workspace with queued assessments. "
           "Pending packages appear with upload date, system name, and status badge.",
           "sca-dashboard"),
      (7,  "ISSM daily portfolio",
           "ISSMs can view all assigned systems' daily logbook status at /issm/daily. "
           "Systems with incomplete daily tasks are highlighted. Rotation day progress "
           "for each ISSO is shown in a compact grid.",
           "issm-daily-portfolio"),
      (8,  "System Owner dashboard",
           "System Owners see their systems with authorization status, upcoming ATO expiry dates, "
           "open risks, and a quick link to submit for authorization review.",
           "system-owner-dash"),
      (9,  "Security Posture dashboard",
           "Navigate to /posture for the aggregate posture view: control coverage by family, "
           "risk heat map, POA&M aging chart, and assessment score trends. "
           "Available to admin, CISO, and AO roles.",
           "posture-dashboard"),
      (10, "Notifications inbox",
           "The bell icon in the top navigation shows unread notification count. Click it or "
           "navigate to /notifications to see all notifications: ATO document status changes, "
           "POA&M milestone alerts, and system assignment notices.",
           "notifications-inbox"),
    ],
    "failures": [
      ("Redirected to wrong dashboard", "Your role assignment may be incorrect. Contact an admin to check your role at /admin/users."),
      ("Missing sidebar links", "Some links are role-gated. If a link you expect is absent, verify your platform role and system assignments."),
      ("Notification count stuck", "Clear browser cache and reload. If persistent, contact the administrator."),
    ],
  },

  # ── 03 VIEW-AS / ROLE SHELL ────────────────────────────────────────────────
  {
    "slug":  "view-as-role-shell",
    "title": "View-As and Role Shell",
    "roles": ["Admin", "ISSM", "SCA"],
    "desc":  (
      "Admins can view the application as any user (View-As) to troubleshoot issues "
      "without sharing credentials. Users with multiple role lenses can switch their "
      "active role view via the Role Shell dropdown, scoping the UI to that role's "
      "perspective without changing their actual account permissions."
    ),
    "prereqs": [
      "Admin account for View-As feature.",
      "Multiple system-role assignments for Role Shell switching.",
    ],
    "steps": [
      (1,  "Open user management",
           "Navigate to /admin/users. The users table lists all accounts with role badges "
           "and status indicators. Find the target user whose perspective you want to inspect.",
           "admin-users-list"),
      (2,  "Activate View-As",
           "Click the View button next to the target user's row. You will be redirected "
           "to the dashboard, now viewing as that user. A persistent orange banner reads "
           "'Viewing as [username] — click to exit'. Your own session cookie is preserved "
           "via an HMAC-signed bsv_user_view cookie.",
           "view-as-banner"),
      (3,  "Navigate in View-As mode",
           "Browse any page as the viewed user. The sidebar, dashboard widgets, and system "
           "access reflect exactly what that user sees. You cannot perform write actions "
           "while in View-As mode — all forms are read-only.",
           "isso-view-dashboard"),
      (4,  "Exit View-As",
           "Click the banner or navigate to /exit-view-as. The bsv_user_view cookie is "
           "cleared and your original admin session is restored.",
           "exit-view-as"),
      (5,  "Switch view (role lens)",
           "Users with multiple roles click the role badge in the top navigation to open "
           "the role switcher dropdown. Select a different role lens. The page reloads "
           "scoped to that role's view without changing your underlying permissions.",
           "switch-view-isso"),
      (6,  "Switch-role-view (admin shell)",
           "Admins can enter a role shell via the admin dropdown → Switch Role. Select a "
           "role name to view the application as if you held that role. A purple 'WORKING AS "
           "[ROLE]' banner appears in the navigation.",
           "switch-role-view"),
      (7,  "Exit role shell",
           "Click the 'Exit Shell' link in the navigation banner or navigate to /exit-shell. "
           "Your admin view is restored.",
           "exit-shell"),
    ],
    "failures": [
      ("View button absent for a user", "The user may be frozen or a reserved account. Frozen accounts cannot be viewed-as."),
      ("Role shell missing a role option", "That role has no active assignments in the system. Role options are dynamically populated."),
      ("Banner persists after exit", "Hard refresh (Ctrl+Shift+R) to clear cached headers. The cookie should have been removed."),
    ],
  },

  # ── 04 SYSTEM LIFECYCLE ────────────────────────────────────────────────────
  {
    "slug":  "system-lifecycle",
    "title": "System Record Lifecycle",
    "roles": ["Admin", "ISSM", "System Owner"],
    "desc":  (
      "A system record in BLACKSITE represents one information system in the RMF program. "
      "It holds the FIPS 199 categorization, boundary description, ATO status, team "
      "assignments, and all associated compliance artifacts. This SOP covers creating, "
      "editing, assigning, archiving, and printing a system record."
    ),
    "prereqs": [
      "Admin or ISSM platform role.",
      "System name, abbreviation, and FIPS categorization ready.",
    ],
    "steps": [
      (1,  "Browse systems list",
           "Navigate to /systems. All active systems are listed with status badges (Authorized, "
           "In Progress, Not Authorized), ATO expiry dates, and assigned ISSO names. Use the "
           "search bar to filter by name or abbreviation.",
           "systems-list"),
      (2,  "Create a new system",
           "Click + New System. Fill in: System Name (required), Abbreviation (2-8 chars), "
           "Boundary Description, FIPS 199 Confidentiality/Integrity/Availability impact levels "
           "(Low/Moderate/High), ePHI flag, EIS flag, and initial ATO notes. "
           "Click Save System.",
           "new-system-form"),
      (3,  "Review system detail",
           "After saving, you land on the system detail page showing all fields and the "
           "tab navigation bar. Verify the Overall Impact was computed correctly (FIPS 199 "
           "max rule). Confirm the system ID and abbreviation are correct before proceeding.",
           "system-detail-overview"),
      (4,  "Edit system record",
           "Click Edit on the system detail page or navigate to /systems/{id}/edit. "
           "Update any field and click Save. Changes are logged in the audit trail. "
           "Do not change the abbreviation after ATO documents have been generated — "
           "the filename convention embeds it.",
           "system-edit-form"),
      (5,  "Manage assignments",
           "Navigate to the Assignments tab or /systems/{id}/assignments. Assign an ISSO, "
           "ISSM, SCA, System Owner, PMO lead, and other duty roles. Each assignment "
           "sends a notification to the assigned user. Remove assignments by clicking "
           "the × next to a user's name.",
           "system-assignments"),
      (6,  "Archive a system",
           "Admin only: navigate to /admin/systems/archived to view archived systems. "
           "To archive an active system, use the Archive action on the system edit page. "
           "Archived systems are hidden from all standard views but retain all data. "
           "Unarchiving restores the system to active status.",
           "archived-systems"),
      (7,  "Print system report",
           "Navigate to /systems/{id}/report or click the Print Report button on the system "
           "detail page. A printer-friendly HTML report opens showing categorization, "
           "team assignments, control summary, and open POA&M list. Use your browser's "
           "Print function (Ctrl+P) to save as PDF.",
           "system-report-print"),
    ],
    "failures": [
      ("Abbreviation already in use", "Each abbreviation must be unique. Choose a different one or check if the system already exists."),
      ("FIPS impact fields blank after save", "Ensure all three CIA fields are set before saving. The form requires all three."),
      ("Assignment not visible to assigned user", "The user may lack a system role assignment. Verify their role at /admin/users."),
    ],
  },

  # ── 05 SYSTEM DETAIL TABS ──────────────────────────────────────────────────
  {
    "slug":  "system-detail-nav",
    "title": "System Detail Tab Navigation",
    "roles": ["All roles assigned to the system"],
    "desc":  (
      "The system detail page organizes all compliance data into tabs. Each tab provides "
      "a focused view of a specific compliance domain. This SOP covers what each tab "
      "contains and when to use it."
    ),
    "prereqs": [
      "Assigned to the system or admin role.",
      "System record exists and is active.",
    ],
    "steps": [
      (1,  "Overview tab",
           "The default tab shows system metadata: name, abbreviation, boundary, FIPS impact "
           "levels, ATO status, expiry date, ePHI/EIS flags, and assigned personnel. "
           "The RMF step progress bar shows where the system stands in the 7-step process.",
           "system-overview-tab"),
      (2,  "Controls tab",
           "Lists all NIST 800-53 Rev 5 controls scoped to this system, with implementation "
           "status, responsible role, and narrative excerpt. Filter by family or status. "
           "Click a control row to open the workspace editor.",
           "controls-tab"),
      (3,  "Inventory tab",
           "Hardware and software inventory for the system boundary. Each item has type, "
           "description, vendor, version, and criticality. Use the + Add Item button to "
           "add new inventory entries.",
           "inventory-tab"),
      (4,  "Connections tab",
           "External and internal connections (ISAs, MOU/As, API integrations). Each record "
           "shows the partner name, data flow direction, and agreement type. Alerts appear "
           "when ISA expiry dates are within 90 days.",
           "connections-tab"),
      (5,  "Artifacts tab",
           "Uploaded compliance artifacts (diagrams, policies, evidence files). Each artifact "
           "has a type label, upload date, and approval status. Artifacts marked 'approved' "
           "are available for ATO document appendix embedding.",
           "artifacts-tab"),
      (6,  "Teams tab",
           "System team roster showing all assigned personnel with their system roles, "
           "duty assignments, and last activity date. BCDR events for the system are "
           "listed in the lower panel.",
           "teams-tab"),
      (7,  "EIS Assessment tab",
           "Enterprise Information System assessment checklist. Mark each criterion as "
           "met/not met with supporting notes. The overall EIS determination (Yes/No/Unknown) "
           "feeds into the system categorization report.",
           "eis-assessment-tab"),
      (8,  "Parameters tab",
           "System-specific parameter overrides for NIST control assignments. Each parameter "
           "maps to a control family. Setting values here auto-populates corresponding "
           "control workspace fields.",
           "parameters-tab"),
      (9,  "Daily Ops tab",
           "Shows today's task completion status (N/M badge) and the current deep work rotation "
           "day. Click Open Daily Hub to access the full daily workflow interface. "
           "This tab is visible to all assigned roles.",
           "daily-ops-tab"),
    ],
    "failures": [
      ("Tab shows 403", "You are not assigned to this system. Ask your ISSO or admin to add you."),
      ("Controls tab empty", "No controls have been added yet. Import from the NIST catalog on the Controls tab."),
      ("Artifacts not showing in ATO appendix", "Artifact must be in 'approved' status. Change status via the artifact detail edit form."),
    ],
  },

  # ── 06 CONTROL CATALOG ────────────────────────────────────────────────────
  {
    "slug":  "control-catalog",
    "title": "NIST Control Catalog and Workspace",
    "roles": ["ISSO", "SCA", "ISSM", "Admin"],
    "desc":  (
      "BLACKSITE maintains a local copy of the NIST SP 800-53 Rev 5 OSCAL catalog. "
      "Controls are browsable, searchable, and can be linked to system records with "
      "implementation narratives, status, and responsible roles. The control workspace "
      "provides a structured editor for each control."
    ),
    "prereqs": [
      "NIST catalog downloaded (automatic on first run, or via /admin/updater).",
      "System record created and assigned.",
    ],
    "steps": [
      (1,  "Browse the catalog",
           "Navigate to /controls. The full NIST 800-53 Rev 5 catalog is displayed in a "
           "searchable table. Filter by control family (AC, AU, CM, etc.) using the family "
           "dropdown. Search by control ID or keyword in the search box.",
           "catalog-browse"),
      (2,  "View control detail — AC-1",
           "Click any control row to open its detail page. The page shows the full control "
           "statement with sub-requirements, supplemental guidance, parameters, and related "
           "controls. Use this page to understand what the control requires before writing "
           "your implementation narrative.",
           "control-detail-ac1"),
      (3,  "View control detail — AC-2",
           "Complex controls like AC-2 (Account Management) have multiple sub-requirements "
           "(a. through j.). Each sub-requirement is listed with its prose. Enhancements "
           "(AC-2(1) through AC-2(13)) appear as nested rows.",
           "control-detail-ac2"),
      (4,  "System controls list",
           "Navigate to the Controls tab on a system detail page. This shows the subset of "
           "controls scoped to the system with implementation status indicators: Complete "
           "(green), Partial (yellow), Insufficient (orange), Not Found (red), NA (gray).",
           "system-controls-list"),
      (5,  "Open the control workspace",
           "Click a control row in the system controls list to open the workspace editor. "
           "The workspace shows the control statement on the left and your implementation "
           "fields on the right: Implementation Status, Responsible Role, and Narrative. "
           "Fill in all three fields for a complete control.",
           "system-control-workspace"),
      (6,  "Import controls from catalog",
           "On the Controls tab, click Import Controls to scope NIST catalog controls to "
           "this system. Select the baseline (Low/Moderate/High) or choose individual "
           "controls. Click Import to add them.",
           "import-controls-page"),
    ],
    "failures": [
      ("Catalog empty at /controls", "The NIST catalog has not been downloaded. Navigate to /admin and trigger a catalog update."),
      ("Workspace save fails", "Check that Implementation Status and Responsible Role are filled. Both are required for PARTIAL or above."),
      ("Import adds no controls", "The selected controls may already be imported. Check the current controls list for existing entries."),
    ],
  },

  # ── 07 ASSESSMENT UPLOAD ──────────────────────────────────────────────────
  {
    "slug":  "assessment-upload",
    "title": "Assessment Package Upload",
    "roles": ["ISSO", "SCA", "Admin"],
    "desc":  (
      "ISSOs submit assessment packages (SSP, supporting evidence, test results) for "
      "SCA review. The upload form accepts PDF and DOCX files up to 20 MB. The AI-assisted "
      "assessor automatically scores control narratives and flags gaps. SCAs then review "
      "the automated findings and provide their professional judgment."
    ),
    "prereqs": [
      "Assessment package prepared as a single PDF or DOCX file (max 20 MB).",
      "System record exists and ISSO is assigned.",
    ],
    "steps": [
      (1,  "Navigate to the upload form",
           "Navigate to /upload. Select the target system from the dropdown. Choose the "
           "assessment file using the file picker. Only PDF and DOCX are accepted. "
           "Click Upload Assessment.",
           "upload-form"),
      (2,  "Monitor processing",
           "After upload, a background task parses the document, extracts control narratives, "
           "and scores each control (0–5 scale). Processing typically takes 30–90 seconds "
           "depending on document size. The admin assessments list at /admin shows the "
           "current status (Pending → Processing → Complete).",
           "admin-assessments-list"),
      (3,  "SCA workspace review",
           "Once processing is complete, the assigned SCA sees the assessment in their "
           "workspace at /sca/workspace. Click the assessment row to open the review view "
           "showing per-control grades, issues, and anomaly flags.",
           "sca-workspace"),
      (4,  "View results",
           "The assessment results page shows: overall quality score, control-by-control grades, "
           "anomaly flags (score anomalies, boilerplate detection, coverage gaps), P1 priority "
           "incomplete controls, and quick wins (controls one element away from PARTIAL).",
           "results-page"),
    ],
    "failures": [
      ("Upload rejected — file too large", "Compress the PDF or split into sections. Maximum file size is 20 MB."),
      ("Processing stuck in Pending", "The background worker may have crashed. Contact the admin to restart the service."),
      ("SCA workspace empty", "The SCA may not be assigned to the system. Check assignments at /systems/{id}/assignments."),
    ],
  },

  # ── 08 SSP EXPORT ─────────────────────────────────────────────────────────
  {
    "slug":  "ssp-export",
    "title": "SSP Generation and Analyzer",
    "roles": ["ISSO", "ISSM", "Admin"],
    "desc":  (
      "BLACKSITE can generate a System Security Plan (SSP) document in two modes: "
      "Controls-only (structured control table) or Full (includes approved artifacts "
      "as appendices). The SSP Analyzer allows admins to upload an external SSP for "
      "automated quality scoring."
    ),
    "prereqs": [
      "System record with controls populated.",
      "Artifacts approved if Full mode appendices are needed.",
    ],
    "steps": [
      (1,  "Select SSP output mode",
           "Navigate to the SSP export section (available from the system detail page or "
           "reports center). The mode selector page shows two options: Controls mode "
           "(control table only) and Full mode (control table + artifact appendices).",
           "ssp-mode-select"),
      (2,  "SSP Analyzer — upload external SSP",
           "Admin only: navigate to /admin/ssp. Click Upload SSP to submit an existing "
           "SSP document for automated scoring. The analyzer parses narratives, scores "
           "quality, and identifies gaps using the same engine as the assessment uploader.",
           "admin-ssp-analyzer"),
      (3,  "Review SSP analysis results",
           "The analysis results page shows per-control grades (Critical/High/Medium/Adequate), "
           "narrative quality indicators, missing element warnings, and a family-level coverage "
           "heat map. Export the findings as a PDF report.",
           "ssp-review-detail"),
      (4,  "Reports center",
           "Navigate to /reports to see all generated reports for all systems. Each row shows "
           "the report type, system, generation date, file size, and a Download button. "
           "Reports are retained until manually deleted.",
           "reports-center"),
    ],
    "failures": [
      ("SSP generation fails with 500", "Check that the system has at least one control with a narrative. Empty systems cannot generate SSPs."),
      ("Full mode appendices missing", "Ensure artifacts are uploaded and their status is set to 'approved' on the Artifacts tab."),
      ("Analyzer stuck in Processing", "Large SSPs (>10 MB) can take up to 5 minutes. If stuck beyond 10 minutes, contact admin."),
    ],
  },

  # ── 09 POAM LIFECYCLE ─────────────────────────────────────────────────────
  {
    "slug":  "poam-lifecycle",
    "title": "POA&M Lifecycle Management",
    "roles": ["ISSO", "ISSM", "SCA", "AO", "Admin"],
    "desc":  (
      "Plan of Action and Milestones (POA&M) items track identified deficiencies "
      "and their remediation plans. BLACKSITE enforces a state machine for POA&M "
      "status transitions with role-based push-power rules. Each status change is "
      "logged in an immutable approval trail."
    ),
    "prereqs": [
      "System record with an assigned ISSO.",
      "Identified finding or vulnerability to track.",
    ],
    "steps": [
      (1,  "View the POA&M list",
           "Navigate to /poam. All POA&M items across your assigned systems are listed "
           "with status badges, severity, control ID, system name, and scheduled completion "
           "date. Use the status filter tabs to narrow the view.",
           "poam-list-all"),
      (2,  "Create a new POA&M item",
           "Click + New POA&M. Fill in: System, Control ID (from NIST catalog datalist), "
           "Title, Weakness Description, Severity (Critical/High/Moderate/Low), Scheduled "
           "Completion Date, Responsible Role, and Remediation Plan. Click Save. "
           "A POA&M ID (e.g., ABVR022826-1001AC01) is auto-generated.",
           "poam-new-form"),
      (3,  "Open POA&M detail (draft status)",
           "Click a POA&M row to open its detail page. In Draft status, the ISSO can edit "
           "all fields. The Approval Trail section at the bottom shows an empty log. "
           "When ready, click Submit for Review to advance to In Progress.",
           "poam-detail-draft"),
      (4,  "Import POA&M items",
           "Navigate to /poam/import to bulk-import POA&M items from CSV or XLSX. "
           "Download the template first, fill in the required columns, then upload. "
           "A preview step shows parsed rows before committing.",
           "poam-import-page"),
      (5,  "Export POA&M list",
           "Navigate to /poam/export to download all open POA&M items as CSV or XLSX. "
           "The export includes all fields, approval trail entries, and evidence filenames. "
           "Use this for FISMA reporting or external review.",
           "poam-export-page"),
      (6,  "In Progress status",
           "When submitted, the POA&M moves to In Progress. The ISSO updates milestone "
           "progress notes and can upload closure evidence. The SCA can review and "
           "request changes or advance to Ready for Review.",
           "poam-status-open"),
      (7,  "Blocked status",
           "If remediation is blocked by an external dependency, the ISSO sets status "
           "to Blocked and fills in the blocker description and responsible party. "
           "Blocked items are flagged in dashboard counts.",
           "poam-blocked-form"),
      (8,  "Ready for Review status",
           "SCA sets the item to Ready for Review after confirming remediation evidence. "
           "The ISSO and SCA both appear in the approval trail. The AO can now "
           "review the item.",
           "poam-ready-for-review"),
      (9,  "Closed / Verified status",
           "AO or admin marks the item Closed after verifying the closure evidence. "
           "A Completion Date must be entered. The approval trail records the closing "
           "official and timestamp. Closed items appear in the audit log.",
           "poam-closed-verified"),
      (10, "Deferred / Waiver status",
           "If remediation cannot be completed by the scheduled date, the ISSO requests "
           "a waiver. Fill in the waiver justification and new target date. "
           "AO approval is required to grant the waiver.",
           "poam-deferred-waiver"),
      (11, "Accepted Risk status",
           "For risks the organization chooses to accept rather than remediate, "
           "the ISSO sets status to Pending AO Acceptance with residual risk justification. "
           "The AO then approves or rejects the risk acceptance.",
           "poam-accepted-risk"),
      (12, "False Positive status",
           "If the finding is determined to be a false positive, the SCA documents "
           "the rationale and sets status to False Positive. No further action is "
           "required. The item remains in the audit trail.",
           "poam-false-positive"),
    ],
    "failures": [
      ("Status button greyed out", "Your role does not have push-power for the target status. Check the POAM_PUSH_POWER role matrix."),
      ("POA&M ID not generated", "IDs are generated on save. If missing after save, check the database connection."),
      ("Completion date field absent", "Completion date only appears for terminal statuses (Closed, Verified). Status must be terminal."),
    ],
  },

  # ── 10 POAM EVIDENCE ──────────────────────────────────────────────────────
  {
    "slug":  "poam-evidence",
    "title": "POA&M Evidence Upload",
    "roles": ["ISSO", "SCA", "Admin"],
    "desc":  (
      "Closure evidence (screenshots, test reports, configuration exports) is attached "
      "directly to POA&M items. Accepted file types are PDF, PNG, JPG, DOCX, XLSX, TXT, "
      "and CSV. Files up to 10 MB per upload. EXE and ZIP files are blocked."
    ),
    "prereqs": [
      "POA&M item exists and is In Progress or later.",
      "Evidence file prepared (PDF, PNG, JPG, DOCX, XLSX, TXT, or CSV).",
    ],
    "steps": [
      (1,  "Open the evidence section",
           "Open the POA&M detail page. Scroll to the Closure Evidence section near the "
           "bottom of the page. The section shows previously uploaded files and an "
           "upload form.",
           "poam-detail-evidence-section"),
      (2,  "Upload evidence file",
           "Click Choose File in the evidence upload form. Select your evidence file "
           "(accepted: PDF, PNG, JPG, DOCX, XLSX, TXT, CSV — max 10 MB). "
           "Add a brief description of what the file demonstrates. Click Upload Evidence. "
           "The file appears in the evidence list with a download link and upload timestamp.",
           "poam-detail-with-evidence"),
    ],
    "failures": [
      ("Upload blocked — EXE or ZIP", "Executable and archive files are blocked for security. Convert content to PDF or screenshot."),
      ("File size error", "Maximum file size is 10 MB. Compress the PDF or split into multiple files."),
      ("Evidence section absent", "The POA&M may still be in Draft status. Submit for review first to enable evidence uploads."),
    ],
  },

  # ── 11 RISK REGISTER ──────────────────────────────────────────────────────
  {
    "slug":  "risk-register",
    "title": "Risk Register Management",
    "roles": ["ISSO", "ISSM", "Admin"],
    "desc":  (
      "The risk register tracks identified risks separate from POA&M items. Risks include "
      "a likelihood × impact score matrix, treatment plan, and status. The heat map on "
      "the risk list provides a visual overview of the risk posture."
    ),
    "prereqs": [
      "System record with ISSO assigned.",
    ],
    "steps": [
      (1,  "View the risk list",
           "Navigate to /risks. All risks across your assigned systems are listed with "
           "risk level badges (Critical/High/Moderate/Low), system name, likelihood, "
           "impact, and treatment status. The risk heat map appears above the table.",
           "risks-list"),
      (2,  "Create a new risk",
           "Click + New Risk. Fill in: System, Title, Description, Likelihood (1–5), "
           "Impact (1–5), Treatment Plan, and Status (Open/Mitigated/Accepted/Closed). "
           "The Risk Score and Level are computed automatically (Likelihood × Impact).",
           "risk-new-form"),
      (3,  "View risk detail",
           "Click a risk row to open its detail page. The detail shows all fields, "
           "the computed risk level, treatment plan, and creation/modification history. "
           "Edit any field and click Save to update.",
           "risk-detail"),
      (4,  "Export risk register",
           "Navigate to /risks/export to download the full risk register as CSV or XLSX. "
           "The export includes computed risk scores, treatment plans, and status. "
           "Use for risk management reporting.",
           "risks-export"),
    ],
    "failures": [
      ("Risk score shows 0", "Both Likelihood and Impact must be set (1–5). Verify both fields are filled."),
      ("Risk level shows Low for high scores", "Risk level is computed as L×I. A score of 15 is Critical. Check if fields saved correctly."),
    ],
  },

  # ── 12 OBSERVATIONS ───────────────────────────────────────────────────────
  {
    "slug":  "observations",
    "title": "Observations and Finding Promotion",
    "roles": ["SCA", "Pen Tester", "ISSO", "Admin"],
    "desc":  (
      "Observations are informal findings noted during assessments or pen tests. "
      "They can be promoted to formal POA&M items when they rise to the level of "
      "a trackable deficiency. This SOP covers creating, reviewing, and promoting observations."
    ),
    "prereqs": [
      "System assigned to user.",
      "Assessment or pen test findings available to document.",
    ],
    "steps": [
      (1,  "View observations list",
           "Navigate to /observations. Observations are listed with type, severity, system, "
           "submitter, and date. Filter by status (Open/Closed) or system using the "
           "filter controls.",
           "observations-list"),
      (2,  "Create a new observation",
           "Click + New Observation. Fill in: System, Type (Technical/Administrative/Physical), "
           "Severity, Title, Description, and Recommended Action. Click Save. "
           "The observation is now visible to the ISSO and SCA for the system.",
           "observation-new-form"),
      (3,  "Review observation detail",
           "Click an observation row to open its detail page. The SCA or ISSO can add "
           "review notes and change the status. If the finding warrants formal tracking, "
           "proceed to promote it to a POA&M item.",
           "observation-detail"),
      (4,  "Promote to POA&M",
           "On the observation detail page, click Promote to POA&M. A pre-filled POA&M "
           "creation form opens with the observation title, description, and system pre-populated. "
           "Complete the remaining POA&M fields (severity, control ID, dates) and save. "
           "The observation status updates to Promoted and links to the new POA&M.",
           "observation-promote"),
    ],
    "failures": [
      ("Promote button absent", "Only SCA and admin roles can promote observations. ISSO can view but not promote."),
      ("Observation not visible to ISSO", "Observations are scoped to the assigned system. Verify the ISSO is assigned."),
    ],
  },

  # ── 13 ATO DOCUMENTS ──────────────────────────────────────────────────────
  {
    "slug":  "ato-documents",
    "title": "ATO Document Package",
    "roles": ["ISSO", "ISSM", "SCA", "AO", "Admin"],
    "desc":  (
      "The Authorization to Operate (ATO) document package contains up to 10 required "
      "artifacts (SSP, FIPS 199, RAR, SAP, SAR, POA&M, ISA, PIA, BCDR, Authorization Decision). "
      "Each document has a defined workflow: Draft → Submitted → Under Review → Approved → Finalized. "
      "The AO makes the final authorization decision from the package."
    ),
    "prereqs": [
      "System record created with FIPS categorization complete.",
      "ISSO, SCA, and AO all assigned.",
    ],
    "steps": [
      (1,  "View the ATO dashboard",
           "Navigate to /ato. The ATO dashboard lists all systems with their authorization status, "
           "ATO expiry date, and number of documents in each status. Pending decisions are "
           "highlighted with action badges.",
           "ato-dashboard"),
      (2,  "Open a system's ATO matrix",
           "Click a system row or navigate to /ato/{system_id}. The matrix table lists all "
           "10 required documents with their current status, last updated date, and available "
           "actions (Generate, Upload, Submit, Approve, Finalize).",
           "ato-system-matrix"),
      (3,  "View document detail",
           "Click a document row (e.g., System Security Plan) to open its detail page. "
           "The page shows document metadata, current status badge, the full status history, "
           "and action buttons appropriate for your role.",
           "ato-doc-detail-ssp"),
      (4,  "Submit a draft document",
           "When the document is ready for review, click Submit. You must confirm the "
           "submission in a dialog. Status changes to Submitted. The SCA and ISSM receive "
           "a notification.",
           "ato-doc-draft-form"),
      (5,  "Generate a document",
           "For auto-generated documents (e.g., FIPS 199 categorization report), click "
           "the Generate button. BLACKSITE builds the document in the background using "
           "system data. When ready, the document appears as a downloadable PDF.",
           "ato-doc-generate"),
      (6,  "Upload a document",
           "For manually prepared documents (e.g., Risk Assessment Report), click Upload "
           "and select the file (PDF or DOCX, max 20 MB). The file is stored and the "
           "status advances to Submitted automatically.",
           "ato-doc-upload"),
      (7,  "Approve a document (SCA/ISSM)",
           "SCA and ISSM can approve documents that are Submitted or Under Review. "
           "Click Approve. The approval is recorded in the status trail with your "
           "username and timestamp.",
           "ato-doc-action-submit"),
      (8,  "Authorized status",
           "Once the AO grants authorization, the system record shows 'Authorized' status "
           "with the authorization date and expiry. The ATO package is locked for the "
           "duration of the authorization period.",
           "ato-doc-approved-status"),
    ],
    "failures": [
      ("Generate button absent", "Not all documents can be auto-generated. Check which documents have a Generate option in the matrix."),
      ("Submit button greyed out", "Document may already be in a non-editable status. Check the current status badge."),
      ("ATO matrix shows wrong documents", "The full 10-document set is shown regardless of system type. NA documents can be marked Not Applicable."),
    ],
  },

  # ── 14 SUBMISSION AND AO DECISION ─────────────────────────────────────────
  {
    "slug":  "submission-ao",
    "title": "Authorization Submission and AO Decision",
    "roles": ["ISSO", "System Owner", "AO", "Admin"],
    "desc":  (
      "When an ATO package is complete, the ISSO submits it for AO review. The AO "
      "reviews the package, makes a risk determination, and records an authorization "
      "decision with duration (1/3/5 years or ongoing). The decision updates the "
      "system's authorization status immediately."
    ),
    "prereqs": [
      "ATO document package complete with all required documents approved.",
      "AO assigned to the system.",
    ],
    "steps": [
      (1,  "Submit for authorization",
           "Navigate to /systems/{id}/submit or click Submit for Authorization on the "
           "system detail page. Confirm the submission in the dialog. Status changes to "
           "Pending Authorization. The AO receives a notification.",
           "submission-form"),
      (2,  "View submissions list",
           "Navigate to /submissions to see all submitted packages with status, system "
           "name, submission date, and AO assignment. ISSOs see only their own systems; "
           "AOs see all pending submissions assigned to them.",
           "submissions-list"),
      (3,  "Submission detail review",
           "Click a submission row to open its detail. The detail page shows the package "
           "contents, all document statuses, open POA&M count, and risk summary. "
           "The AO uses this view to make their determination.",
           "submission-detail"),
      (4,  "AO decisions queue",
           "AOs navigate to /ao/decisions to see all packages awaiting their decision. "
           "Each row shows the system name, ISSO, submission date, and outstanding "
           "risk items. Click a row to open the decision form.",
           "ao-decisions-list"),
      (5,  "Record AO decision",
           "On the AO decision form, select: Authorize (select duration: 1yr/3yr/5yr/ "
           "ongoing/custom), Deny, or Authorize with Conditions. Add decision notes. "
           "Click Record Decision. The system's auth_status and expiry date are updated "
           "immediately. The ISSO receives a notification.",
           "ao-decision-form"),
      (6,  "Authorized status confirmed",
           "Return to the system detail page. The overview tab shows 'Authorized' status "
           "badge, authorization date, and expiry date. The ATO package is now locked "
           "for the authorization period.",
           "auth-status-authorized"),
    ],
    "failures": [
      ("Submit button unavailable", "Required ATO documents may not all be in Approved status. Check the ATO matrix."),
      ("AO decision form missing duration picker", "The AO must select an authorization type (Authorize) before the duration pills appear."),
      ("Status still Pending after decision", "Refresh the page. If still pending after 30 seconds, check the audit log for errors."),
    ],
  },

  # ── 15 REPORTS AND EXPORTS ────────────────────────────────────────────────
  {
    "slug":  "reports-exports",
    "title": "Reports and Data Exports",
    "roles": ["ISSO", "ISSM", "SCA", "Admin"],
    "desc":  (
      "BLACKSITE provides multiple report and export options: system-level PDF reports, "
      "generated compliance reports (executive summary, evidence pack, BCDR pack), "
      "POA&M exports (CSV/XLSX), risk exports, and audit log exports."
    ),
    "prereqs": [
      "Reports are generated on demand or auto-triggered on rotation days 18/24/25.",
    ],
    "steps": [
      (1,  "Reports center",
           "Navigate to /reports for an overview of all generated reports across all systems. "
           "Filter by system, report type, or date range. Each report shows file size, "
           "generation date, and status (generating/ready/error).",
           "reports-center-main"),
      (2,  "System print report",
           "Navigate to /systems/{id}/report. This page renders a printer-optimized HTML "
           "report of the system's current state. Use Ctrl+P to save as PDF for offline "
           "distribution or archival.",
           "system-report-print"),
      (3,  "System generated reports list",
           "Navigate to /systems/{id}/reports to see all PDF reports generated for this "
           "system. Each row has a Download button that serves the file with the proper "
           "Content-Disposition header for browser save.",
           "system-reports-list"),
      (4,  "Export POA&M list",
           "Navigate to /poam/export. Select format (CSV or XLSX) and optional date range "
           "filter. Click Export. The file downloads immediately.",
           "poam-export"),
      (5,  "Export risk register",
           "Navigate to /risks/export. Same format and filter options as POA&M export. "
           "The export includes computed risk scores and treatment plans.",
           "risks-export"),
      (6,  "Audit log export",
           "Admin only: navigate to /admin/audit. Click Export Audit Log to download the "
           "full audit trail as CSV. The audit log records every significant action with "
           "timestamp, user, action type, and detail string.",
           "audit-export"),
    ],
    "failures": [
      ("Report status stuck on 'generating'", "The background task may have failed. Check the server logs. Trigger a new generation."),
      ("Download returns 404", "The report file may have been deleted from disk. Regenerate the report."),
      ("POA&M export empty", "Verify your user has assigned systems with POA&M items. The export is scoped to your assignments."),
    ],
  },

  # ── 16 USER MANAGEMENT ────────────────────────────────────────────────────
  {
    "slug":  "user-management",
    "title": "User Account Management",
    "roles": ["Admin"],
    "desc":  (
      "Admins manage all user accounts: creating, provisioning, freezing, removing, "
      "and role assignment. BLACKSITE uses a reservation system to hold usernames for "
      "provisioned accounts. Shadow users are accounts that have never logged in."
    ),
    "prereqs": [
      "Admin account.",
    ],
    "steps": [
      (1,  "View users list",
           "Navigate to /admin/users. All accounts are listed with platform role badges, "
           "last login date, status (active/frozen/shadow), and action buttons. "
           "The search bar filters by username or display name.",
           "admin-users-list"),
      (2,  "Add a user account",
           "Click + Add User. Enter the username (must match the Authelia/proxy username "
           "exactly), Display Name, and Platform Role. Click Add User. "
           "The account is created and the user can log in immediately.",
           "add-user-form"),
      (3,  "Provision a user (guided setup)",
           "Navigate to /admin/users/provision. Provisioning creates a user with a "
           "pre-configured role and optional system assignment in one step. "
           "Fill in username, display name, role, and optionally a system to assign. "
           "Click Provision. A notification email is sent if mail is configured.",
           "provision-form"),
      (4,  "Freeze a user account",
           "Click Freeze next to a user. Confirm in the dialog. Frozen accounts cannot "
           "log in and their sessions are invalidated immediately. The account remains "
           "in the database. Freeze is reversible via the Unfreeze action.",
           "user-freeze-confirm"),
      (5,  "Remove a user account",
           "Click Remove next to a user. Confirm in the dialog. The username is held in "
           "RemovedUserReservation for 1 year (prevents username reuse). All assignments "
           "for the user are removed. This action is logged in the audit trail.",
           "user-remove-confirm"),
      (6,  "View name reservations",
           "Navigate to /admin/users/reservations to see all held usernames from removed "
           "accounts. Each reservation shows the original account's removal date and "
           "1-year expiry. Admins can override and release a reservation early if needed.",
           "user-reservations"),
      (7,  "Shadow users",
           "Navigate to /admin/users/shadow to see accounts that have been provisioned "
           "but have never logged in. These accounts may represent onboarding delays or "
           "incorrect usernames. Follow up with the user or remove if stale.",
           "shadow-users"),
      (8,  "Bulk role assignment",
           "On the users list, use the role selector dropdown next to each user to change "
           "their platform role individually. Changes take effect on the user's next page load. "
           "Role changes are logged in the audit trail.",
           "bulk-role-assign"),
      (9,  "Configure max packages per user",
           "ISSMs can set a per-user maximum concurrent assessment package limit on the "
           "ISSM dashboard. This controls how many simultaneous packages an ISSO can have "
           "in the assessment queue.",
           "max-packages-setting"),
    ],
    "failures": [
      ("Add User fails — username exists", "The username is already taken or reserved. Check the reservations list."),
      ("Freeze has no effect on active session", "Sessions expire after the 15-minute idle timeout. The freeze blocks future logins immediately."),
      ("Provision email not sent", "Mail may not be configured. Check /admin/system-settings for SMTP configuration."),
    ],
  },

  # ── 17 AUDIT LOG ──────────────────────────────────────────────────────────
  {
    "slug":  "audit-log",
    "title": "Audit Log Review",
    "roles": ["Admin"],
    "desc":  (
      "BLACKSITE maintains a comprehensive audit log of all significant actions: "
      "logins, role changes, POA&M updates, ATO decisions, file uploads, and admin actions. "
      "The audit log is append-only and cannot be modified by any user."
    ),
    "prereqs": [
      "Admin account.",
    ],
    "steps": [
      (1,  "View the audit log",
           "Navigate to /admin/audit. The log shows the most recent 500 entries by default. "
           "Each entry shows timestamp, username, action type, and detail string. "
           "Entries are color-coded by category (auth=blue, write=green, admin=orange, error=red).",
           "audit-log-list"),
      (2,  "Filter by user or action",
           "Use the filter bar at the top to narrow by username, action type, or date range. "
           "Click Apply Filters. The table updates to show only matching entries.",
           "audit-log-filter-user"),
      (3,  "Export audit log",
           "Click Export CSV at the top of the audit log page. The full filtered audit "
           "log downloads as a CSV file with all columns. Use this for compliance reporting "
           "or incident investigation.",
           "audit-log-export-link"),
    ],
    "failures": [
      ("Audit log empty", "No actions have been logged yet, or the filter is too narrow. Clear filters and reload."),
      ("Export file empty", "Verify the date range includes events. Audit log entries older than the configured retention period may be purged."),
    ],
  },

  # ── 18 SIEM EVENTS ─────────────────────────────────────────────────────────
  {
    "slug":  "siem-events",
    "title": "SIEM Security Events",
    "roles": ["Admin", "ISSO"],
    "desc":  (
      "BLACKSITE receives and displays SIEM events via the /api/siem/ingest API endpoint. "
      "Events are categorized by severity (Critical/High/Medium/Low/Info) and type. "
      "ISSOs and admins can review events, filter by severity, and drill into event details."
    ),
    "prereqs": [
      "SIEM integration configured and sending events to /api/siem/ingest.",
    ],
    "steps": [
      (1,  "View SIEM event list",
           "Navigate to /admin/siem. Events are listed in reverse chronological order "
           "with severity badge, event type, source IP, timestamp, and brief description. "
           "The total event count appears at the top.",
           "siem-list"),
      (2,  "Filter by severity",
           "Click a severity filter button (Critical, High, Medium, Low, Info) to narrow "
           "the view. Multiple filters can be active simultaneously. The count badge "
           "updates to reflect filtered results.",
           "siem-filter-severity"),
      (3,  "Drill into event details",
           "Click an event row to expand its detail panel. The detail shows the full "
           "event payload, source system, affected host, and raw event data. "
           "Use the Create POA&M link to promote a significant event to a tracked finding.",
           "siem-drilldown"),
    ],
    "failures": [
      ("No events showing", "SIEM integration may not be configured, or no events have been ingested. Check integration settings."),
      ("Create POA&M link absent", "This action requires ISSO or admin role. Verify your role assignment."),
    ],
  },

  # ── 19 BCDR WORKFLOW ──────────────────────────────────────────────────────
  {
    "slug":  "bcdr-workflow",
    "title": "BCDR Event and Team Management",
    "roles": ["BCDR Coordinator", "ISSO", "Admin"],
    "desc":  (
      "Business Continuity and Disaster Recovery (BCDR) events track activations, "
      "exercises, and recovery tests. BCDR Coordinators create events, document the "
      "timeline, and record sign-offs from stakeholders. System teams are managed "
      "separately as the organizational roster for each system."
    ),
    "prereqs": [
      "BCDR Coordinator role assigned.",
      "System record with teams tab populated.",
    ],
    "steps": [
      (1,  "View BCDR dashboard",
           "Navigate to /bcdr/dashboard. Active and past BCDR events are listed with "
           "type (Activation/Exercise/Test), start date, status, and involved systems. "
           "Click + New Event to create one.",
           "bcdr-dashboard"),
      (2,  "Create a BCDR event",
           "Click + New Event. Fill in: Event Type, Affected Systems (multi-select), "
           "Description, Start Time, and Estimated Duration. Click Create. "
           "The event opens in Active status. All assigned BCDR personnel receive a notification.",
           "bcdr-create-event"),
      (3,  "Update event detail",
           "Click an active event to open its detail. Update the timeline notes as the "
           "event progresses. Document recovery actions taken, systems restored, and "
           "any issues encountered.",
           "bcdr-event-detail"),
      (4,  "Record stakeholder sign-off",
           "When the event is resolved, click Record Sign-off. Enter the stakeholder "
           "name and role. Add final resolution notes. Click Complete Event. "
           "The event status changes to Closed and a completion timestamp is recorded.",
           "bcdr-signoff"),
      (5,  "Manage system teams",
           "Navigate to /systems/{id}/teams to view and edit the system team roster. "
           "Team members are listed with their roles and contact information. "
           "Use this page to keep the emergency contact list current.",
           "system-teams-page"),
    ],
    "failures": [
      ("Create event fails — no systems listed", "You must be assigned to at least one system as BCDR Coordinator."),
      ("Sign-off button absent", "Only the event creator and admin can record sign-offs. Verify the event owner."),
    ],
  },

  # ── 20 DAILY OPS ──────────────────────────────────────────────────────────
  {
    "slug":  "daily-ops",
    "title": "Daily Operations Hub",
    "roles": ["ISSO", "ISSM", "SCA", "System Owner", "PMO", "Pen Tester", "Auditor",
              "Incident Responder", "BCDR Coordinator", "Data Owner"],
    "desc":  (
      "The Daily Operations Hub provides a structured daily workflow for each assigned role. "
      "Each role sees a customized set of 2–8 daily tasks. Completing tasks and saving "
      "the logbook creates an audit trail of daily activities. Sub-forms capture structured "
      "data for specific tasks (change review, backup check, access spot-check)."
    ),
    "prereqs": [
      "Assigned to the system.",
      "System record active.",
    ],
    "steps": [
      (1,  "Open the Daily Hub (ISSO view)",
           "Navigate to /systems/{id}/daily. The ISSO sees all 8 task cards with checkboxes. "
           "The system name and today's date are pre-populated. Metric snapshot (open POA&Ms, "
           "risks, observations, incidents) appears in the header strip.",
           "daily-hub-isso"),
      (2,  "Complete and save the daily logbook",
           "Check each task as you complete it. Add optional notes in the Notes field. "
           "Click Save Daily Log. The logbook row is written to the database with the "
           "task flags and metric snapshot. The Daily Ops tab badge updates to show N/8.",
           "daily-save-form"),
      (3,  "View logbook history",
           "Navigate to /systems/{id}/daily/history. A 30-day calendar grid shows each "
           "day's completion status (green = all tasks complete, yellow = partial, "
           "gray = no log). Click any day to view that day's logbook.",
           "daily-history"),
      (4,  "Change review sub-form (Task 2)",
           "Navigate to /systems/{id}/daily/change-review. Fill in: ticket references, "
           "high-risk change count, all-approved flag, backout plan exists flag, and "
           "any untracked changes found. Click Submit. A ChangeReviewRecord is created.",
           "change-review-form"),
      (5,  "Backup check sub-form (Task 4)",
           "Navigate to /systems/{id}/daily/backup-check. Fill in: result (pass/fail/partial), "
           "affected ePHI systems, job health status, and whether an issue was raised. "
           "Click Submit. A BackupCheckRecord is created.",
           "backup-check-form"),
      (6,  "Access spot-check sub-form (Task 5 — HIPAA)",
           "Navigate to /systems/{id}/daily/access-spotcheck. Fill in: records sampled count, "
           "anomaly found flag, terminated user found flag, and anomaly description if applicable. "
           "Click Submit. An AccessSpotCheck record is created.",
           "access-spotcheck-form"),
      (7,  "Daily Hub (ISSM view)",
           "ISSMs see tasks 1, 2, 6, and 8. The layout is identical to the ISSO view "
           "but scoped to ISSM responsibilities. Task labels are customized per role.",
           "daily-hub-issm"),
      (8,  "Daily Hub (SCA view)",
           "SCAs see tasks 3, 6, and 8. The SCA-specific task labels appear in place "
           "of ISSO task labels. All other functionality is identical.",
           "daily-hub-sca"),
    ],
    "failures": [
      ("Daily Hub shows 403", "You must be assigned to the system. Contact the ISSO to add your role assignment."),
      ("Task checkboxes missing", "Your role may not have any tasks configured. Contact the admin to verify ROLE_TASK_CONFIGS."),
      ("Save fails", "Ensure the system_id in the URL is valid. A 422 error indicates a form validation issue."),
    ],
  },

  # ── 21 DEEP WORK ROTATION ─────────────────────────────────────────────────
  {
    "slug":  "deep-work-rotation",
    "title": "Deep Work Rotation",
    "roles": ["ISSO", "ISSM", "SCA", "System Owner", "PMO", "Pen Tester", "Auditor",
              "Incident Responder", "BCDR Coordinator", "Data Owner"],
    "desc":  (
      "The Deep Work Rotation is a structured 25-day cycle of 90-minute focused work sessions. "
      "Each day has a specific compliance topic with instructions and required evidence. "
      "Days 18, 24, and 25 auto-generate PDF reports on completion. Quarterly overrides "
      "replace specific rotation slots with time-bound activities."
    ),
    "prereqs": [
      "Assigned to the system.",
      "Daily Hub active (rotation initializes on first visit).",
    ],
    "steps": [
      (1,  "View current rotation day",
           "Navigate to /systems/{id}/rotation. The current day's content appears with: "
           "day number (1–25), topic title, estimated duration, detailed instructions, "
           "and the evidence label describing what to submit. Click Start Session if "
           "using the 90-minute timer.",
           "rotation-current-day"),
      (2,  "View rotation history",
           "Navigate to /systems/{id}/rotation/history. Past completion records are listed "
           "with day number, completion date, notes, and evidence filename. "
           "Click a row to view details.",
           "rotation-history"),
      (3,  "View rotation calendar",
           "Navigate to /systems/{id}/rotation/calendar. The 25-day grid shows each day "
           "with its topic, completion status (green = done, empty = pending), and federal "
           "holiday indicators (gray background). The current day is highlighted.",
           "rotation-calendar"),
      (4,  "Complete a rotation day",
           "On the current day page, fill in the evidence notes and optionally upload an "
           "evidence file. Click Complete Day. The rotation day advances by 1 (wrapping "
           "from 25 to 1). A DeepWorkCompletion record is created. On days 18, 24, and 25, "
           "a PDF report generation is triggered automatically.",
           "rotation-complete-form"),
      (5,  "ISSM portfolio rotation view",
           "ISSMs navigate to /issm/daily to see all assigned ISSOs' rotation progress. "
           "Each ISSO's current rotation day and last completion date are shown. "
           "ISSOs who have not completed a rotation day this week are highlighted.",
           "issm-daily-portfolio"),
    ],
    "failures": [
      ("Rotation stuck on day 25", "The completion record may be missing. Check /rotation/history for the Day 25 entry. If absent, manually complete."),
      ("Auto-generated report not appearing", "Check /systems/{id}/reports. The BackgroundTask may still be running. Wait 2 minutes and refresh."),
      ("Quarterly override not showing", "Quarterly overrides only appear when the current date falls in the override window. Check QUARTERLY_OVERRIDES config."),
    ],
  },

  # ── 22 COMPLIANCE RECORDS ─────────────────────────────────────────────────
  {
    "slug":  "compliance-records",
    "title": "Compliance Record Management",
    "roles": ["ISSO", "ISSM", "Admin"],
    "desc":  (
      "BLACKSITE tracks five compliance record types per system: Vendors/BAAs, "
      "Interconnection Records (ISAs), Data Flow Records, Privacy Assessments (PTA/PIA), "
      "and Restore Test Records. These records provide the evidence base for continuous "
      "monitoring and ATO maintenance."
    ),
    "prereqs": [
      "System record with ISSO assigned.",
    ],
    "steps": [
      (1,  "Vendor and BAA registry",
           "Navigate to /systems/{id}/vendors. The vendor list shows all third-party vendors "
           "with service type, ePHI flag, BAA status, BAA expiry date, and contact. "
           "Click + Add Vendor to add a new entry. BAAs expiring within 90 days are "
           "highlighted in yellow.",
           "vendors-list"),
      (2,  "Interconnection records (ISAs)",
           "Navigate to /systems/{id}/interconnections. Each ISA record shows the partner "
           "system, data types exchanged, ISA existence/expiry, monitoring status, and "
           "encryption indicators. Add or update records as new connections are established.",
           "interconnections-list"),
      (3,  "Data flow records",
           "Navigate to /systems/{id}/dataflows. Data flows document API integrations and "
           "data pipelines: authentication method, transit/rest encryption, logging status, "
           "and termination procedures. Required for HIPAA technical safeguard documentation.",
           "dataflows-list"),
      (4,  "Privacy assessments (PTA/PIA)",
           "Navigate to /systems/{id}/privacy-assessments. PTAs (Privacy Threshold Analysis) "
           "and PIAs (Privacy Impact Assessment) are listed by assess type. Each record "
           "tracks data elements collected, purpose, disclosures, retention policy, and "
           "review status (draft/current/needs_review).",
           "privacy-assessments-list"),
      (5,  "Restore test records",
           "Navigate to /systems/{id}/restore-tests. Backup restore tests are documented "
           "with test date, scope, result (pass/fail/partial), time-to-restore in minutes, "
           "validator name, and optional evidence file. HIPAA requires annual testing.",
           "restore-tests-list"),
    ],
    "failures": [
      ("Vendor BAA expiry not alerting", "Alerts appear only for BAAs expiring within 90 days. Check the BAA expiry date field is set correctly."),
      ("PIA status stuck on draft", "Status must be manually changed to 'current' after completing the assessment. Edit the record and update status."),
      ("Restore test result not saving", "All required fields (test_date, scope, result) must be filled. Check for validation errors."),
    ],
  },

  # ── 23 EMPLOYEE QUIZ ──────────────────────────────────────────────────────
  {
    "slug":  "employee-quiz",
    "title": "Daily Employee Security Quiz",
    "roles": ["All roles"],
    "desc":  (
      "BLACKSITE presents a daily security awareness question on the dashboard. "
      "Answering correctly extends your streak. Quiz scores feed into the combined "
      "assessment score (30% weight). Streaks and scores are visible on the dashboard "
      "and contribute to the Allstar designation."
    ),
    "prereqs": [
      "Authenticated session.",
      "Quiz for today not yet completed.",
    ],
    "steps": [
      (1,  "Quiz prompt on dashboard",
           "The quiz card appears on your dashboard if today's quiz is not yet answered. "
           "The card shows the question category and a 'Take Quiz' button. "
           "Dismissed questions cannot be answered later today.",
           "quiz-dashboard-card"),
      (2,  "Answer the quiz question",
           "Click Take Quiz or navigate to /dashboard/quiz. The question appears with "
           "4 answer options. Select your answer and click Submit.",
           "quiz-page"),
      (3,  "Submit and see result",
           "After submitting, the correct answer is highlighted. If correct, a success "
           "message shows your current streak. If incorrect, the correct answer explanation "
           "appears. Your quiz score is updated in the database.",
           "quiz-submit"),
      (4,  "Streak and score on dashboard",
           "Return to the dashboard. The quiz widget now shows your current streak count "
           "and today's result (correct/incorrect). The streak resets to 0 if you miss "
           "a day or answer incorrectly.",
           "quiz-streak-result"),
    ],
    "failures": [
      ("Quiz card not showing", "You may have already answered today's quiz. It resets at midnight UTC."),
      ("Quiz page shows 'No question today'", "The quiz database may not have a question for today. Contact admin to load quiz questions."),
    ],
  },

  # ── 24 PROFILE ─────────────────────────────────────────────────────────────
  {
    "slug":  "profile-preferences",
    "title": "Profile and Display Preferences",
    "roles": ["All roles"],
    "desc":  (
      "Users can customize their BLACKSITE experience: display name, font size, "
      "data density, rows per page, theme, and RSS feed subscriptions. "
      "Preferences are stored server-side and persist across devices."
    ),
    "prereqs": [
      "Authenticated session.",
    ],
    "steps": [
      (1,  "Open profile page",
           "Click your username in the top navigation and select Profile, or navigate "
           "to /profile. The profile page shows your account details and a preferences "
           "panel for display settings.",
           "profile-page"),
      (2,  "RSS feed subscriptions",
           "Navigate to /profile/feeds or click the Feeds link on the profile page. "
           "A curated list of 15 cybersecurity RSS feeds is available. Toggle subscriptions "
           "on/off. Subscribed feeds appear in the News section of your dashboard.",
           "profile-feeds"),
      (3,  "Theme switcher",
           "Click your username dropdown and select a theme from the theme swatches section. "
           "14 themes are available: default, dark, terminal, synthwave, crimson, ocean, "
           "amber, noir, and more. Theme is stored in a cookie and applied on page load.",
           "theme-switcher"),
    ],
    "failures": [
      ("Theme not persisting", "Ensure cookies are enabled for the BLACKSITE domain. The theme is stored in the bsv_theme cookie."),
      ("Feed not appearing on dashboard", "After subscribing, reload the dashboard. Feed items may take 15 minutes to populate from the remote RSS source."),
    ],
  },

  # ── 25 ADMIN SETTINGS ─────────────────────────────────────────────────────
  {
    "slug":  "admin-settings",
    "title": "Admin System Settings",
    "roles": ["Admin"],
    "desc":  (
      "System-wide configuration is managed through the Admin System Settings page. "
      "Settings include: chat enable/disable, autofail configuration, RSS feed source "
      "management, and data ingestion. These settings take effect immediately without "
      "a service restart."
    ),
    "prereqs": [
      "Admin account.",
    ],
    "steps": [
      (1,  "System settings page",
           "Navigate to /admin/system-settings. The settings page shows key-value pairs "
           "for all configurable settings. Current values and descriptions are shown. "
           "Click the toggle or edit field for each setting and save.",
           "system-settings-page"),
      (2,  "RSS feed source management",
           "Navigate to /admin/feeds. The 15 curated feed sources are listed with title, "
           "URL, and category. Admins can add new feed sources or remove existing ones. "
           "Changes affect all users' subscription options.",
           "feed-sources-admin"),
      (3,  "Autofail configuration",
           "Navigate to /admin/autofail. Configure the autofail threshold: minimum quiz "
           "score and minimum combined score required to pass assessments. "
           "ISSOs below threshold are automatically flagged in the ISSM dashboard.",
           "autofail-config"),
      (4,  "Data ingestion",
           "Navigate to /admin/ingest. Upload a CSV, XLSX, or JSON file to bulk-import "
           "users or systems. Download the template first. A preview step shows parsed rows "
           "before committing. Ingestion logs appear in the audit trail.",
           "data-ingest"),
    ],
    "failures": [
      ("Settings not taking effect", "Settings are cached in memory for 5 minutes. Wait and retry, or restart the service."),
      ("Ingest preview shows 0 rows", "Verify the file format matches the template. The first row must be a header row matching expected column names."),
    ],
  },

  # ── 26 GLOBAL SEARCH ──────────────────────────────────────────────────────
  {
    "slug":  "global-search",
    "title": "Global Search",
    "roles": ["All roles"],
    "desc":  (
      "The global search at /search finds systems, controls, POA&M items, risks, "
      "observations, and users across all data you have access to. Results are "
      "grouped by type and ranked by relevance."
    ),
    "prereqs": [
      "Authenticated session.",
    ],
    "steps": [
      (1,  "Open the search page",
           "Navigate to /search or press the search icon in the top navigation. "
           "The search page loads with an empty results state and a search input field.",
           "search-empty"),
      (2,  "Search for a term",
           "Type a search term and press Enter or click Search. Results appear grouped "
           "by type (Systems, Controls, POA&Ms, Risks). Each result shows the item name, "
           "type badge, and a link to the item's detail page.",
           "search-results"),
      (3,  "Search by control ID",
           "Type a NIST control ID (e.g., 'ac-2') to find the control directly. "
           "The result links to both the catalog detail page and any system-control "
           "workspace where that control is implemented.",
           "search-suggest"),
    ],
    "failures": [
      ("No results for known item", "Search is scoped to your role's accessible data. If you lack system access, those systems' data will not appear."),
      ("Search page returns 404", "The /search route may not be enabled. Contact the admin."),
    ],
  },

]

# ─── Document styling helpers ─────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_cover(doc: Document, title: str, roles: list[str]):
    """Add a cover section to the document."""
    # Top rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("━" * 70)
    run.font.color.rgb = BRAND_CYAN
    run.font.size = Pt(7)

    # System label
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{SYSTEM_NAME}  ·  {SYSTEM_ID}")
    run.font.color.rgb = BRAND_GRAY
    run.font.size = Pt(9)

    # Title
    p = doc.add_heading(title, level=1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = TEXT_DARK
        run.font.size = Pt(20)

    # Meta row
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Roles:  {', '.join(roles)}")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_GRAY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Version:  {DOC_VERSION}    |    Date:  {TODAY}    |    Classification:  INTERNAL USE ONLY")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_GRAY

    # Bottom rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(16)
    run = p.add_run("━" * 70)
    run.font.color.rgb = BRAND_CYAN
    run.font.size = Pt(7)

def add_prereqs(doc: Document, prereqs: list[str]):
    if not prereqs:
        return
    h = doc.add_heading("Prerequisites", level=2)
    h.paragraph_format.space_before = Pt(6)
    for run in h.runs:
        run.font.color.rgb = TEXT_DARK
        run.font.size = Pt(12)
    for item in prereqs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10)

def add_description(doc: Document, desc: str):
    h = doc.add_heading("Overview", level=2)
    h.paragraph_format.space_before = Pt(12)
    for run in h.runs:
        run.font.color.rgb = TEXT_DARK
        run.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(desc)
    run.font.size = Pt(10)

def add_step(doc: Document, slug: str, step_num: int, heading: str,
             body: str, screenshot_label: str | None):
    """Add a numbered step with optional screenshot."""
    # Step heading
    h = doc.add_heading(f"Step {step_num}: {heading}", level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(11)

    # Body text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(body)
    run.font.size = Pt(10)

    # Screenshot
    if screenshot_label:
        img_path = SHOT_DIR / f"{slug}-step-{step_num:02d}-{screenshot_label}.png"
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(6.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_para.paragraph_format.space_after = Pt(4)
                # Caption
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                run = cap.add_run(f"Figure {step_num}: {heading}")
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = BRAND_GRAY
            except Exception as e:
                p = doc.add_paragraph()
                run = p.add_run(f"[Screenshot unavailable: {e}]")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = BRAND_GRAY
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"[Screenshot not found: {img_path.name}]")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = BRAND_GRAY

def add_failures(doc: Document, failures: list[tuple[str, str]]):
    if not failures:
        return
    h = doc.add_heading("Common Issues and Resolutions", level=2)
    h.paragraph_format.space_before = Pt(16)
    for run in h.runs:
        run.font.color.rgb = TEXT_DARK
        run.font.size = Pt(12)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = tbl.rows[0].cells
    hdr[0].text = "Symptom"
    hdr[1].text = "Resolution"
    for cell in hdr:
        set_cell_bg(cell, "1a1a2e")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = BRAND_WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    for symptom, resolution in failures:
        row = tbl.add_row().cells
        row[0].text = symptom
        row[1].text = resolution
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    tbl.columns[0].width = Inches(2.5)
    tbl.columns[1].width = Inches(4.0)

def build_sop(wf: dict) -> Path:
    """Build a single workflow SOP DOCX."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    add_cover(doc, wf["title"], wf["roles"])
    add_description(doc, wf["desc"])
    add_prereqs(doc, wf.get("prereqs", []))

    # Steps heading
    h = doc.add_heading("Procedure", level=2)
    h.paragraph_format.space_before = Pt(16)
    for run in h.runs:
        run.font.color.rgb = TEXT_DARK
        run.font.size = Pt(12)

    for step_num, heading, body, screenshot_label in wf["steps"]:
        add_step(doc, wf["slug"], step_num, heading, body, screenshot_label)

    add_failures(doc, wf.get("failures", []))

    out_path = OUT_DIR / f"SOP-{wf['slug'].upper()}.docx"
    doc.save(str(out_path))
    return out_path

# ─── Index document ───────────────────────────────────────────────────────────

def build_index(workflow_paths: list[tuple[str, str, Path]]) -> Path:
    """Build the master index DOCX."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    add_cover(doc, "Help Center SOP Index", ["All Roles"])

    p = doc.add_paragraph()
    run = p.add_run(
        f"This index lists all Standard Operating Procedure documents for the {SYSTEM_NAME}. "
        f"Each SOP provides step-by-step instructions with screenshots for a specific workflow. "
        f"Generated {TODAY}. Document version {DOC_VERSION}."
    )
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(16)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["#", "Workflow", "Roles", "Filename"]):
        hdr[i].text = label
        set_cell_bg(hdr[i], "1a1a2e")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = BRAND_WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    for n, (title, roles, path) in enumerate(workflow_paths, 1):
        row = tbl.add_row().cells
        row[0].text = str(n)
        row[1].text = title
        row[2].text = roles
        row[3].text = path.name
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    tbl.columns[0].width = Inches(0.35)
    tbl.columns[1].width = Inches(2.8)
    tbl.columns[2].width = Inches(2.2)
    tbl.columns[3].width = Inches(2.2)

    out_path = OUT_DIR / "SOP-INDEX.docx"
    doc.save(str(out_path))
    return out_path

# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Generating {len(WORKFLOWS)} SOP documents to {OUT_DIR}")
    index_entries = []
    for wf in WORKFLOWS:
        path = build_sop(wf)
        index_entries.append((wf["title"], ", ".join(wf["roles"][:2]) + ("..." if len(wf["roles"]) > 2 else ""), path))
        step_count = len(wf["steps"])
        print(f"  [ok]  {path.name}  ({step_count} steps)")

    idx = build_index(index_entries)
    print(f"  [ok]  {idx.name}  (index)")

    # Summary
    docx_files = list(OUT_DIR.glob("SOP-*.docx"))
    total_size = sum(f.stat().st_size for f in docx_files)
    print(f"\nDone. {len(docx_files)} documents, {total_size // 1024} KB total.")
